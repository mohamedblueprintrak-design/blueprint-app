import os
import io
import logging
import base64
import json
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import SessionLocal, Project, Analysis, BOQItem, Defect, UploadedFile, SiteVisit, SiteVisitImage

logger = logging.getLogger("pdf_generator")

# محاولة تحميل الخط العربي، وإلا استخدام الخط الافتراضي
FONT_PATH = "fonts/Amiri-Regular.ttf"
FONT_NAME = "Amiri"
try:
    if os.path.exists(FONT_PATH):
        pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
    else:
        logger.warning("Font file not found, falling back to Helvetica.")
        FONT_NAME = "Helvetica"
except Exception as e:
    logger.error(f"Error loading font: {e}")
    FONT_NAME = "Helvetica"

def clean_text_for_pdf(text):
    """استبدال الرموز التعبيرية بنصوص بديلة لتجنب ظهور أحرف غريبة في PDF"""
    if not isinstance(text, str):
        text = str(text)
    replacements = {
        "💰": "اقتصادي",
        "🪶": "خفيف",
        "🛡️": "آمن",
        "🔴": "عالية",
        "🟡": "متوسطة",
        "🟢": "منخفضة",
        "✅": "تم الحل",
        "⏳": "مفتوح",
        "📊": "نتائج",
        "⚠️": "تحذير",
        "🔍": "بحث",
        "📐": "هندسي",
        "📏": "قياس",
        "🧱": "خرسانة",
        "🏗️": "إنشاء",
        "🪚": "نجارة",
        "🔩": "حديد",
        "📄": "مستند",
        "📎": "ملف",
        "📅": "تاريخ",
        "📍": "موقع",
        "📚": "مرجع",
        "💬": "محادثة",
        "🤖": "ذكاء اصطناعي",
        "🪄": "Blue",
    }
    for emoji, replacement in replacements.items():
        text = text.replace(emoji, replacement)
    return text

def draw_arabic_text(c, text, x, y, font_size=12):
    """
    دالة مساعدة لرسم النصوص العربية بشكل صحيح في PDF
    """
    try:
        # تنظيف النص من الرموز التعبيرية أولاً
        cleaned_text = clean_text_for_pdf(text)
        # إعادة تشكيل النص العربي
        reshaped_text = arabic_reshaper.reshape(cleaned_text)
        # ضبط اتجاه النص
        bidi_text = get_display(reshaped_text)
        c.setFont(FONT_NAME, font_size)
        c.drawString(x, y, bidi_text)
    except Exception as e:
        # في حالة الفشل، استخدم الخط العادي
        c.setFont("Helvetica", font_size)
        c.drawString(x, y, text)

def generate_project_report(project_id):
    """إنشاء تقرير PDF للمشروع"""
    db = SessionLocal()
    try:
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            return None

        # جلب البيانات
        analyses = db.query(Analysis).filter(Analysis.project_id == project_id).order_by(Analysis.created_at.desc()).all()
        boqs = db.query(BOQItem).filter(BOQItem.project_id == project_id).all()
        defects = db.query(Defect).filter(Defect.project_id == project_id).all()
        site_visits = db.query(SiteVisit).filter(SiteVisit.project_id == project_id).order_by(SiteVisit.visit_date.desc()).limit(5).all()

        # إنشاء ملف PDF
        filename = f"report_{project_id}.pdf"
        c = canvas.Canvas(filename, pagesize=A4)
        w, h = A4

        # تاريخ التقرير
        draw_arabic_text(c, f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 50, h - 20, 10)

        # العنوان الرئيسي
        draw_arabic_text(c, f"تقرير مشروع: {project.name}", 50, h - 50, 24)
        draw_arabic_text(c, f"الموقع: {project.location}", 50, h - 80, 12)
        c.line(50, h - 90, w - 50, h - 90)

        y = h - 120

        # ---------- التحليلات الهندسية الهامة ----------
        draw_arabic_text(c, "التحليلات الهندسية الهامة", 50, y, 16)
        y -= 30

        # تصفية التحليلات: نأخذ فقط الـ 10 تحليلات الأخيرة التي ليست مجرد "chat"
        important_analyses = [a for a in analyses if a.task_type != "chat"][:10]

        for ana in important_analyses:
            if y < 100:
                c.showPage()
                y = h - 50
                draw_arabic_text(c, "التحليلات الهندسية الهامة (تابع)", 50, y, 16)
                y -= 30

            # اسم المهمة
            draw_arabic_text(c, f"- {ana.task_type}: {ana.input_text[:50]}...", 50, y, 12)
            y -= 20

            # محتوى النتيجة
            res_data = ana.result_json.get("results", {}) if ana.result_json else {}
            text_content = res_data.get("📊 النتائج", "") or res_data.get("📄 تحليل الملف", "")
            if text_content:
                short_text = str(text_content).replace("**", "")[:100]
                draw_arabic_text(c, short_text, 70, y, 10)
                y -= 20

            # رسم الكمرة (صورة) إذا كانت موجودة
            img_data = res_data.get("image_data")
            if img_data:
                try:
                    img_bytes = base64.b64decode(img_data)
                    img = ImageReader(io.BytesIO(img_bytes))
                    if y - 120 < 50:
                        c.showPage()
                        y = h - 50
                    c.drawImage(img, 50, y - 120, width=150, height=100, preserveAspectRatio=True)
                    y -= 140
                except Exception as e:
                    logger.error(f"فشل رسم الصورة: {e}")
                    draw_arabic_text(c, "[تعذر رسم الصورة]", 70, y, 10)
                    y -= 20

            # بيانات DXF (اختياري)
            dxf_data = res_data.get("dxf_data")
            if dxf_data:
                draw_arabic_text(c, "(يوجد ملف DXF للتصميم)", 70, y, 10)
                y -= 20

            y -= 10

        # إضافة ملاحظة بعدد المحادثات العادية المحذوفة
        chat_count = len([a for a in analyses if a.task_type == "chat"])
        if chat_count > 0:
            draw_arabic_text(c, f"(تم تجاهل {chat_count} محادثة عادية)", 70, y, 10)
            y -= 20

        # ---------- العيوب المسجلة ----------
        if y < 150:
            c.showPage()
            y = h - 50
            draw_arabic_text(c, "العيوب المسجلة", 50, y, 16)
            y -= 30
        else:
            y -= 20
            draw_arabic_text(c, "العيوب المسجلة", 50, y, 16)
            y -= 30

        for defect in defects:
            if y < 100:
                c.showPage()
                y = h - 50
                draw_arabic_text(c, "العيوب المسجلة (تابع)", 50, y, 16)
                y -= 30
            line = f"- {defect.description} | {defect.severity} | {defect.status}"
            draw_arabic_text(c, line, 50, y, 10)
            y -= 20
            if defect.image_id:
                img_file = db.query(UploadedFile).filter(UploadedFile.id == defect.image_id).first()
                if img_file and os.path.exists(img_file.file_path):
                    try:
                        c.drawImage(img_file.file_path, 50, y-80, width=100, height=80, preserveAspectRatio=True)
                        y -= 100
                    except:
                        pass
            y -= 10

        # ---------- آخر زيارات الموقع ----------
        if y < 150:
            c.showPage()
            y = h - 50
            draw_arabic_text(c, "آخر زيارات الموقع", 50, y, 16)
            y -= 30
        else:
            y -= 20
            draw_arabic_text(c, "آخر زيارات الموقع", 50, y, 16)
            y -= 30

        for visit in site_visits:
            if y < 100:
                c.showPage()
                y = h - 50
                draw_arabic_text(c, "آخر زيارات الموقع (تابع)", 50, y, 16)
                y -= 30
            line = f"- {visit.visit_date.strftime('%Y-%m-%d')}: {visit.location_name or 'بدون موقع'}"
            draw_arabic_text(c, line, 50, y, 10)
            y -= 20
            first_img = db.query(SiteVisitImage).filter(SiteVisitImage.site_visit_id == visit.id).first()
            if first_img and os.path.exists(first_img.image_path):
                try:
                    c.drawImage(first_img.image_path, 50, y-80, width=120, height=80, preserveAspectRatio=True)
                    y -= 100
                except:
                    pass
            y -= 10

        # ---------- جدول الكميات (BOQ) ----------
        if y < 150:
            c.showPage()
            y = h - 50
            draw_arabic_text(c, "جدول الكميات (BOQ)", 50, y, 16)
            y -= 30
        else:
            y -= 20
            draw_arabic_text(c, "جدول الكميات (BOQ)", 50, y, 16)
            y -= 30

        # تجميع البنود المتشابهة
        grouped_items = defaultdict(lambda: {"quantity": 0, "total_price": 0, "unit": ""})

        for item in boqs:
            key = item.description
            grouped_items[key]["quantity"] += item.quantity
            grouped_items[key]["total_price"] += item.total_price
            grouped_items[key]["unit"] = item.unit

        total_cost = 0
        for desc, values in grouped_items.items():
            price = values["total_price"]
            total_cost += price
            line = f"{desc} | {values['quantity']} {values['unit']} | {price} جنيه"
            draw_arabic_text(c, line, 50, y, 10)
            y -= 20

        y -= 10
        draw_arabic_text(c, f"الإجمالي التقديري: {total_cost} جنيه", 50, y, 14)

        c.save()
        with open(filename, "rb") as f:
            pdf_bytes = f.read()
        os.remove(filename)
        return pdf_bytes

    except Exception as e:
        logger.error(f"PDF Gen Error: {e}")
        import traceback
        traceback.print_exc()
        return None
    finally:
        db.close()

def generate_project_word(project_id):
    """إنشاء تقرير Word للمشروع (قابل للتعديل)"""
    db = SessionLocal()
    try:
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            return None

        # جلب البيانات
        analyses = db.query(Analysis).filter(Analysis.project_id == project_id).order_by(Analysis.created_at.desc()).all()
        boqs = db.query(BOQItem).filter(BOQItem.project_id == project_id).all()
        defects = db.query(Defect).filter(Defect.project_id == project_id).all()
        site_visits = db.query(SiteVisit).filter(SiteVisit.project_id == project_id).order_by(SiteVisit.visit_date.desc()).limit(5).all()

        # إنشاء مستند Word
        doc = Document()

        # إعداد اتجاه المستند من اليمين إلى اليسار (للعربية)
        try:
            doc.sections[0].start_type
        except:
            pass

        # العنوان الرئيسي
        title = doc.add_heading(clean_text_for_pdf(f"تقرير مشروع: {project.name}"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(clean_text_for_pdf(f"الموقع: {project.location}"))
        doc.add_paragraph(clean_text_for_pdf(f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M')}"))
        doc.add_paragraph("_" * 50)

        # ---------- التحليلات الهندسية الهامة ----------
        doc.add_heading("التحليلات الهندسية الهامة", level=1)
        important_analyses = [a for a in analyses if a.task_type != "chat"][:10]

        for ana in important_analyses:
            doc.add_paragraph(clean_text_for_pdf(f"مهمة: {ana.task_type}"), style="List Bullet")
            doc.add_paragraph(clean_text_for_pdf(f"السؤال: {ana.input_text[:100]}..."), style="List Bullet 2")
            
            res_data = ana.result_json.get("results", {}) if ana.result_json else {}
            text_content = res_data.get("📊 النتائج", "") or res_data.get("📄 تحليل الملف", "")
            if text_content:
                doc.add_paragraph(clean_text_for_pdf(f"النتيجة: {text_content[:200]}..."), style="List Bullet 2")

        chat_count = len([a for a in analyses if a.task_type == "chat"])
        if chat_count > 0:
            doc.add_paragraph(clean_text_for_pdf(f"(تم تجاهل {chat_count} محادثة عادية)"), style="Italic")

        # ---------- العيوب المسجلة ----------
        doc.add_heading("العيوب المسجلة", level=1)
        if defects:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = clean_text_for_pdf("الوصف")
            hdr_cells[1].text = clean_text_for_pdf("الشدة")
            hdr_cells[2].text = clean_text_for_pdf("الحالة")
            for defect in defects:
                row_cells = table.add_row().cells
                row_cells[0].text = clean_text_for_pdf(defect.description)
                row_cells[1].text = clean_text_for_pdf(defect.severity)
                row_cells[2].text = clean_text_for_pdf(defect.status)
        else:
            doc.add_paragraph("لا توجد عيوب مسجلة.")

        # ---------- آخر زيارات الموقع ----------
        doc.add_heading("آخر زيارات الموقع", level=1)
        if site_visits:
            for visit in site_visits:
                doc.add_paragraph(clean_text_for_pdf(f"- {visit.visit_date.strftime('%Y-%m-%d')}: {visit.location_name or 'بدون موقع'}"), style="List Bullet")
        else:
            doc.add_paragraph("لا توجد زيارات مسجلة.")

        # ---------- جدول الكميات (BOQ) ----------
        doc.add_heading("جدول الكميات (BOQ)", level=1)
        if boqs:
            # تجميع البنود المتشابهة
            grouped_items = defaultdict(lambda: {"quantity": 0, "total_price": 0, "unit": ""})
            for item in boqs:
                key = item.description
                grouped_items[key]["quantity"] += item.quantity
                grouped_items[key]["total_price"] += item.total_price
                grouped_items[key]["unit"] = item.unit

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = clean_text_for_pdf("الوصف")
            hdr_cells[1].text = clean_text_for_pdf("الكمية")
            hdr_cells[2].text = clean_text_for_pdf("السعر (جنيه)")
            
            total_cost = 0
            for desc, values in grouped_items.items():
                row_cells = table.add_row().cells
                row_cells[0].text = clean_text_for_pdf(desc)
                row_cells[1].text = clean_text_for_pdf(f"{values['quantity']} {values['unit']}")
                price = values["total_price"]
                total_cost += price
                row_cells[2].text = clean_text_for_pdf(f"{price} جنيه")
            
            doc.add_paragraph(clean_text_for_pdf(f"الإجمالي التقديري: {total_cost} جنيه"), style="Strong")
        else:
            doc.add_paragraph("لا توجد بنود حصر.")

        # حفظ الملف
        filename = f"report_{project_id}.docx"
        doc.save(filename)
        with open(filename, "rb") as f:
            word_bytes = f.read()
        os.remove(filename)
        return word_bytes

    except Exception as e:
        logger.error(f"Word Gen Error: {e}")
        import traceback
        traceback.print_exc()
        return None
    finally:
        db.close()